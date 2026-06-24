"""Render dei documenti AUA: conversione Markdown -> DOCX e rendering dei template Jinja2.

Due responsabilita':
1. ``markdown_to_docx`` — mini-convertitore da un sottoinsieme di Markdown a ``.docx``
   tramite ``python-docx`` (titoli, paragrafi, elenchi, tabelle, grassetto inline).
2. ``renderizza_template`` — rende un template ``*.md.j2`` con Jinja2, restituendo
   il Markdown gia' popolato con i dati del contesto.
"""

from __future__ import annotations

import pathlib
import re
from typing import Any

import docx
from docx.document import Document as TipoDocumento
from docx.text.paragraph import Paragraph
from jinja2 import ChainableUndefined, Environment, FileSystemLoader, select_autoescape

# Cartella che contiene i template *.md.j2.
CARTELLA_TEMPLATE = pathlib.Path(__file__).parent / "templates"

# Riconosce i frammenti in grassetto inline: **testo**.
_RE_GRASSETTO = re.compile(r"\*\*(.+?)\*\*")


def _aggiungi_testo_con_grassetto(paragrafo: Paragraph, testo: str) -> None:
    """Aggiunge ``testo`` al paragrafo interpretando il grassetto inline ``**...**``.

    Spezza la stringa nei segmenti delimitati da ``**`` e applica il grassetto
    soltanto ai segmenti racchiusi tra i doppi asterischi.
    """
    posizione = 0
    for match in _RE_GRASSETTO.finditer(testo):
        # Testo normale che precede il frammento in grassetto.
        if match.start() > posizione:
            paragrafo.add_run(testo[posizione:match.start()])
        # Frammento in grassetto (senza gli asterischi).
        run = paragrafo.add_run(match.group(1))
        run.bold = True
        posizione = match.end()
    # Eventuale coda di testo normale dopo l'ultimo frammento in grassetto.
    if posizione < len(testo):
        paragrafo.add_run(testo[posizione:])


def _e_riga_tabella(riga: str) -> bool:
    """Vero se la riga e' una riga di tabella Markdown (inizia e finisce con '|')."""
    spogliata = riga.strip()
    return spogliata.startswith("|") and spogliata.endswith("|")


def _e_riga_separatrice(riga: str) -> bool:
    """Vero se la riga e' la riga separatrice di una tabella, tipo ``|---|---|``."""
    spogliata = riga.strip().strip("|")
    # Ogni cella deve contenere solo trattini, due punti e spazi.
    celle = spogliata.split("|")
    return all(re.fullmatch(r"[\s:-]+", cella) and "-" in cella for cella in celle)


def _celle_riga(riga: str) -> list[str]:
    """Estrae le celle da una riga di tabella Markdown."""
    # Si tolgono i delimitatori esterni e si divide sui '|' interni.
    contenuto = riga.strip().strip("|")
    return [cella.strip() for cella in contenuto.split("|")]


def _aggiungi_tabella(documento: TipoDocumento, righe_tabella: list[str]) -> None:
    """Costruisce una tabella DOCX da un blocco di righe Markdown.

    La prima riga e' l'intestazione (in grassetto); l'eventuale riga separatrice
    e' gia' stata esclusa dal chiamante.
    """
    if not righe_tabella:
        return

    matrice = [_celle_riga(riga) for riga in righe_tabella]
    n_colonne = max(len(r) for r in matrice)

    tabella = documento.add_table(rows=0, cols=n_colonne)
    tabella.style = "Table Grid"

    for indice_riga, celle in enumerate(matrice):
        riga_docx = tabella.add_row()
        for indice_colonna in range(n_colonne):
            testo = celle[indice_colonna] if indice_colonna < len(celle) else ""
            cella_docx = riga_docx.cells[indice_colonna]
            # Si usa il paragrafo gia' presente nella cella.
            paragrafo = cella_docx.paragraphs[0]
            if indice_riga == 0:
                # Intestazione: testo in grassetto.
                run = paragrafo.add_run(testo)
                run.bold = True
            else:
                _aggiungi_testo_con_grassetto(paragrafo, testo)


def markdown_to_docx(testo: str, percorso: str) -> None:
    """Converte un sottoinsieme di Markdown in un file ``.docx``.

    Elementi gestiti:
    - titoli ``#``, ``##``, ``###`` (livelli 1-3);
    - paragrafi di testo;
    - elenchi puntati con prefisso ``- ``;
    - tabelle ``|...|`` (la riga separatrice ``|---|`` viene saltata);
    - grassetto inline ``**...**``.
    """
    documento = docx.Document()
    righe = testo.splitlines()

    indice = 0
    while indice < len(righe):
        riga = righe[indice]
        spogliata = riga.strip()

        # Riga vuota: nessun elemento.
        if not spogliata:
            indice += 1
            continue

        # Tabella: si accumulano tutte le righe contigue di tabella.
        if _e_riga_tabella(riga):
            blocco: list[str] = []
            while indice < len(righe) and _e_riga_tabella(righe[indice]):
                if not _e_riga_separatrice(righe[indice]):
                    blocco.append(righe[indice])
                indice += 1
            _aggiungi_tabella(documento, blocco)
            continue

        # Titoli.
        if spogliata.startswith("### "):
            documento.add_heading(spogliata[4:].strip(), level=3)
            indice += 1
            continue
        if spogliata.startswith("## "):
            documento.add_heading(spogliata[3:].strip(), level=2)
            indice += 1
            continue
        if spogliata.startswith("# "):
            documento.add_heading(spogliata[2:].strip(), level=1)
            indice += 1
            continue

        # Elenco puntato.
        if spogliata.startswith("- "):
            paragrafo = documento.add_paragraph(style="List Bullet")
            _aggiungi_testo_con_grassetto(paragrafo, spogliata[2:].strip())
            indice += 1
            continue

        # Paragrafo normale.
        paragrafo = documento.add_paragraph()
        _aggiungi_testo_con_grassetto(paragrafo, spogliata)
        indice += 1

    documento.save(percorso)


def renderizza_template(nome_template: str, contesto: dict[str, Any]) -> str:
    """Rende il template ``nome_template`` (in ``aua/templates``) con il contesto dato.

    Restituisce il Markdown popolato come stringa. L'autoescape e' disattivato
    perche' l'output e' Markdown (non HTML).
    """
    ambiente = Environment(
        loader=FileSystemLoader(str(CARTELLA_TEMPLATE)),
        autoescape=select_autoescape(enabled_extensions=(), default=False),
        # ChainableUndefined permette l'accesso ad attributi mancanti (es. impianto.comune
        # quando 'impianto' non esiste) senza errori: resta Undefined e il filtro
        # 'default' dei template lo sostituisce con i segnaposto [DA COMPLETARE: ...].
        undefined=ChainableUndefined,
        trim_blocks=True,
        lstrip_blocks=True,
        keep_trailing_newline=True,
    )
    template = ambiente.get_template(nome_template)
    return template.render(**contesto)
