#!/usr/bin/env python3
"""Generatore di pratiche AUA (Autorizzazione Unica Ambientale) — Regione Campania.

DPR 13 marzo 2013, n. 59 — VERSIONE A FILE UNICO (autosufficiente).

Tutto è incluso in questo singolo file: lettura/validazione della checklist,
motore di applicabilità dei titoli, template documentali e generazione dei Word.
Non servono altri file del progetto: basta questo .py più le tre librerie
(python-docx, Jinja2, PyYAML).

USO RAPIDO
----------
    # 1) crea una checklist vuota da compilare
    python aua_campania.py crea-checklist mia_checklist.yaml

    # 1-bis) oppure scrivi un esempio già compilato (officina meccanica)
    python aua_campania.py esempio esempio.yaml

    # 2) verifica quali titoli AUA servono
    python aua_campania.py applicabilita mia_checklist.yaml

    # 3) genera il pacchetto Word della pratica
    python aua_campania.py genera mia_checklist.yaml -o output/cliente

INSTALLAZIONE DIPENDENZE
------------------------
    pip install python-docx Jinja2 PyYAML

AVVERTENZA
----------
Produce BOZZE tecniche a supporto del professionista. I template sono scheletri
con segnaposto [DA COMPLETARE], non la modulistica ufficiale: prima di ogni
presentazione reale al SUAP va verificata la modulistica AUA vigente di Regione
Campania / Provincia competente.
"""

from __future__ import annotations

import argparse
import pathlib
import re
import sys
from typing import Any

try:
    import yaml
    import docx
    from docx.text.paragraph import Paragraph
    from jinja2 import (
        ChainableUndefined,
        DictLoader,
        Environment,
        TemplateError,
        TemplateNotFound,
        select_autoescape,
    )
except ImportError as exc:  # pragma: no cover - dipende dall'ambiente
    nome = getattr(exc, "name", "sconosciuta")
    sys.stderr.write(
        f"Dipendenza mancante: {nome}\n"
        "Installa i pacchetti necessari con:\n"
        "    pip install python-docx Jinja2 PyYAML\n"
    )
    raise SystemExit(1)


__version__ = "0.1.0"


# =========================================================================== #
#  TEMPLATE DOCUMENTALI (Markdown + Jinja2, incorporati)                      #
# =========================================================================== #

TEMPLATE_MODELLO_UNICO = """# Istanza di Autorizzazione Unica Ambientale (AUA)

**Ai sensi del DPR 13 marzo 2013, n. 59**

---

## Allo Sportello Unico per le Attività Produttive (SUAP)

Comune di {{ impianto.comune | default("[DA COMPLETARE: Comune competente]") }}
({{ impianto.provincia | default("[DA COMPLETARE: provincia]") }})

Per il tramite del SUAP, all'Autorità competente al rilascio
(Provincia / Città Metropolitana competente per territorio).

**Marca da bollo:** [DA COMPLETARE: estremi marca da bollo da € 16,00 sull'istanza]

---

## 1. Dati del gestore (richiedente)

- **Ragione sociale:** {{ gestore.ragione_sociale | default("[DA COMPLETARE: ragione sociale]") }}
- **Forma giuridica:** {{ gestore.forma_giuridica | default("[DA COMPLETARE: forma giuridica]") }}
- **Codice fiscale:** {{ gestore.codice_fiscale | default("[DA COMPLETARE: codice fiscale]") }}
- **Partita IVA:** {{ gestore.partita_iva | default("[DA COMPLETARE: partita IVA]") }}

### Sede legale

{% set sede = gestore.sede_legale | default({}) %}
- **Indirizzo:** {{ sede.indirizzo | default("[DA COMPLETARE: indirizzo sede legale]") }}
- **Comune:** {{ sede.comune | default("[DA COMPLETARE: comune]") }}
- **Provincia:** {{ sede.provincia | default("[DA COMPLETARE: provincia]") }}
- **CAP:** {{ sede.cap | default("[DA COMPLETARE: CAP]") }}

## 2. Legale rappresentante

{% set lr = gestore.legale_rappresentante | default({}) %}
- **Nome:** {{ lr.nome | default("[DA COMPLETARE: nome]") }}
- **Cognome:** {{ lr.cognome | default("[DA COMPLETARE: cognome]") }}
- **Codice fiscale:** {{ lr.codice_fiscale | default("[DA COMPLETARE: codice fiscale]") }}
- **PEC:** {{ lr.pec | default("[DA COMPLETARE: indirizzo PEC per le comunicazioni]") }}

in qualità di legale rappresentante dell'impresa sopra indicata,

## 3. Dati dell'impianto / insediamento produttivo

- **Denominazione:** {{ impianto.denominazione | default("[DA COMPLETARE: denominazione impianto]") }}
- **Indirizzo:** {{ impianto.indirizzo | default("[DA COMPLETARE: indirizzo impianto]") }}
- **Comune:** {{ impianto.comune | default("[DA COMPLETARE: comune]") }}
- **Provincia:** {{ impianto.provincia | default("[DA COMPLETARE: provincia]") }}
- **CAP:** {{ impianto.cap | default("[DA COMPLETARE: CAP]") }}

### Dati catastali

- **Foglio:** {{ impianto.foglio | default("[DA COMPLETARE: foglio catastale]") }}
- **Particella:** {{ impianto.particella | default("[DA COMPLETARE: particella]") }}

### Inquadramento dell'attività

- **Codice ATECO:** {{ impianto.codice_ateco | default("[DA COMPLETARE: codice ATECO]") }}
- **Descrizione attività:** {{ impianto.descrizione_attivita | default("[DA COMPLETARE: descrizione sintetica del ciclo produttivo]") }}
- **Superficie (mq):** {{ impianto.superficie_mq | default("[DA COMPLETARE: superficie in mq]") }}
- **Numero addetti:** {{ impianto.numero_addetti | default("[DA COMPLETARE: numero addetti]") }}

## 4. Titoli ambientali richiesti con la presente AUA

L'Autorizzazione Unica Ambientale qui richiesta sostituisce e coordina i seguenti
titoli ambientali (art. 3 DPR 59/2013):

{% if titoli %}
| Cod. | Titolo | Riferimento normativo |
|------|--------|-----------------------|
{% for t in titoli %}
| {{ t.codice | default("-") }} | {{ t.nome | default("[DA COMPLETARE: titolo]") }} | {{ t.riferimento | default("[DA COMPLETARE: riferimento]") }} |
{% endfor %}
{% else %}
[DA COMPLETARE: elenco dei titoli ambientali richiesti — nessun titolo determinato dal motore di applicabilità]
{% endif %}

## 5. Dichiarazioni sostitutive (DPR 28 dicembre 2000, n. 445)

Il sottoscritto legale rappresentante, **consapevole delle sanzioni penali**
previste dall'art. 76 del DPR 445/2000 in caso di dichiarazioni mendaci e della
decadenza dai benefici (art. 75) conseguente a dichiarazioni non veritiere, ai
sensi degli artt. 46 e 47 del medesimo decreto, sotto la propria responsabilità

**DICHIARA:**

- che i dati e le informazioni contenuti nella presente istanza e nei relativi
  allegati corrispondono al vero;
- che l'impianto **non è soggetto** alla disciplina dell'Autorizzazione Integrata
  Ambientale (AIA) di cui al Titolo III-bis della Parte Seconda del D.Lgs. 152/2006;
- di impegnarsi a comunicare ogni variazione dei dati dichiarati e a rispettare
  le prescrizioni dell'atto autorizzativo che verrà rilasciato;
- [DA COMPLETARE: ulteriori dichiarazioni richieste dalla modulistica AUA vigente
  di Regione Campania / Provincia competente].

## 6. Durata dell'autorizzazione

L'Autorizzazione Unica Ambientale ha **durata di 15 (quindici) anni** dalla data
di rilascio, ai sensi dell'art. 3, comma 6, del DPR 59/2013.

## 7. Allegati

- [DA COMPLETARE: elenco degli allegati tecnici e amministrativi richiesti
  (relazioni tecniche, planimetrie, copia documento d'identità, ricevuta versamenti)].

---

## 8. Sottoscrizione

Luogo e data: [DA COMPLETARE: luogo e data]

Il Legale Rappresentante
{{ lr.nome | default("") }} {{ lr.cognome | default("") }}

**Firma digitale del legale rappresentante**
(documento sottoscritto digitalmente ai sensi del D.Lgs. 82/2005 — CAD)

[Spazio per la firma digitale]
"""

TEMPLATE_EMISSIONI = """# Relazione tecnica — Emissioni in atmosfera

**Titolo C / D dell'AUA — DPR 59/2013**

{% set em = emissioni_atmosfera | default({}) %}

## 1. Dati dell'impianto

- **Denominazione:** {{ impianto.denominazione | default("[DA COMPLETARE: denominazione impianto]") }}
- **Comune:** {{ impianto.comune | default("[DA COMPLETARE: comune]") }} ({{ impianto.provincia | default("[DA COMPLETARE: provincia]") }})
- **Attività:** {{ impianto.descrizione_attivita | default("[DA COMPLETARE: descrizione del ciclo produttivo]") }}

## 2. Regime autorizzativo

{% if em.regime == "ordinaria" %}
Le emissioni in atmosfera sono assoggettate ad **autorizzazione in via ordinaria**
ai sensi dell'**art. 269 del D.Lgs. 152/2006**.
{% elif em.regime == "generale" %}
Le emissioni in atmosfera sono assoggettate ad **autorizzazione di carattere generale**
ai sensi dell'**art. 272, comma 2, del D.Lgs. 152/2006**.

- **Attività in deroga (Allegato richiamato dall'art. 272):** {{ em.attivita_in_deroga | default("[DA COMPLETARE: attività dell'Allegato — Parte II — richiamata]") }}
{% else %}
[DA COMPLETARE: regime autorizzativo — "ordinaria" (art. 269) oppure "generale" (art. 272 c. 2)]
{% endif %}

## 3. Descrizione del ciclo produttivo e delle fasi che generano emissioni

[DA COMPLETARE: descrizione del ciclo produttivo, delle fasi di lavorazione e
delle relative emissioni in atmosfera (convogliate e/o diffuse)]

## 4. Punti di emissione

{% set punti = em.punti_emissione | default([]) %}
{% if punti %}
| Sigla | Provenienza | Inquinanti | Portata (Nmc/h) | Altezza (m) | Impianto di abbattimento |
|-------|-------------|------------|-----------------|-------------|--------------------------|
{% for p in punti %}
| {{ p.sigla | default("[DA COMPLETARE]") }} | {{ p.provenienza | default("[DA COMPLETARE]") }} | {{ (p.inquinanti | join(", ")) if p.inquinanti else "[DA COMPLETARE]" }} | {{ p.portata_nmc_h | default("[DA COMPLETARE]") }} | {{ p.altezza_m | default("[DA COMPLETARE]") }} | {{ p.impianto_abbattimento | default("[DA COMPLETARE]") }} |
{% endfor %}
{% else %}
[DA COMPLETARE: tabella dei punti di emissione (sigla, provenienza, inquinanti,
portata, altezza, impianto di abbattimento)]
{% endif %}

## 5. Valori limite di emissione di riferimento

I valori limite di emissione applicabili sono quelli fissati:

- per il regime **ordinario (art. 269)**, dall'autorizzazione, sulla base dei
  valori limite e delle prescrizioni dell'**Allegato I alla Parte Quinta** del
  D.Lgs. 152/2006 e dei piani e programmi di qualità dell'aria;
- per il regime **generale (art. 272 c. 2)**, dall'Allegato tecnico
  dell'autorizzazione generale adottata dall'Autorità competente, riferito alla
  specifica attività.

[DA COMPLETARE: prospetto dei valori limite di emissione applicabili per ciascun
inquinante e relativi metodi di campionamento/analisi]

---

*Relazione redatta a supporto dell'istanza AUA. I valori limite e le prescrizioni
vanno verificati rispetto alla normativa e alla modulistica vigenti.*
"""

TEMPLATE_SCARICHI = """# Relazione tecnica — Scarichi di acque reflue

**Titolo A dell'AUA — DPR 59/2013 — art. 124 D.Lgs. 152/2006**

{% set sc = scarichi | default({}) %}

## 1. Dati dell'impianto

- **Denominazione:** {{ impianto.denominazione | default("[DA COMPLETARE: denominazione impianto]") }}
- **Comune:** {{ impianto.comune | default("[DA COMPLETARE: comune]") }} ({{ impianto.provincia | default("[DA COMPLETARE: provincia]") }})

## 2. Tipologia degli scarichi

{% set tipologie = sc.tipologie | default([]) %}
{% if tipologie %}
Le acque reflue prodotte dall'insediamento sono delle seguenti tipologie:

{% for tipo in tipologie %}
- {{ tipo }}
{% endfor %}
{% else %}
[DA COMPLETARE: tipologie di reflui — industriali, domestiche, meteoriche di dilavamento]
{% endif %}

- **Portata complessiva (mc/anno):** {{ sc.portata_mc_anno | default("[DA COMPLETARE: portata in mc/anno]") }}
- **Presenza di sostanze pericolose (Tab. 3/A o 5, All. 5, Parte III):** {{ "Sì" if sc.presenza_sostanze_pericolose else "No" }}

## 3. Schema di trattamento

[DA COMPLETARE: descrizione dello schema di trattamento (linea acque, presidi
depurativi, eventuale pretrattamento) prima dello scarico]

## 4. Recapito e corpo recettore

{% if sc.recapito == "fognatura" %}
Lo scarico recapita in **pubblica fognatura**.
{% elif sc.recapito == "acque_superficiali" %}
Lo scarico recapita in **acque superficiali**.
{% elif sc.recapito == "suolo" %}
Lo scarico recapita sul **suolo**.
{% else %}
[DA COMPLETARE: recapito dello scarico — fognatura / acque superficiali / suolo]
{% endif %}

- **Corpo recettore / gestore fognatura:** {{ sc.corpo_recettore | default("[DA COMPLETARE: denominazione corpo idrico o gestore della fognatura]") }}

## 5. Punti di scarico

{% set punti = sc.punti_scarico | default([]) %}
{% if punti %}
| Sigla | Origine | Trattamento |
|-------|---------|-------------|
{% for p in punti %}
| {{ p.sigla | default("[DA COMPLETARE]") }} | {{ p.origine | default("[DA COMPLETARE]") }} | {{ p.trattamento | default("[DA COMPLETARE]") }} |
{% endfor %}
{% else %}
[DA COMPLETARE: tabella dei punti di scarico (sigla, origine, trattamento)]
{% endif %}

## 6. Valori limite di emissione

Lo scarico deve rispettare i **valori limite di emissione** di cui alla
**Tabella 3 dell'Allegato 5 alla Parte Terza del D.Lgs. 152/2006**, in funzione
del recapito (acque superficiali o pubblica fognatura). Per lo scarico sul suolo
si applicano i limiti della Tabella 4 del medesimo Allegato.

[DA COMPLETARE: prospetto dei parametri da monitorare con i relativi valori limite
applicabili e i metodi analitici di riferimento]

---

*Relazione redatta a supporto dell'istanza AUA. I limiti e le prescrizioni vanno
verificati rispetto alla normativa e alla modulistica vigenti.*
"""

TEMPLATE_RIFIUTI = """# Relazione tecnica — Recupero di rifiuti in procedura semplificata

**Titolo G dell'AUA — DPR 59/2013 — artt. 215-216 D.Lgs. 152/2006**

{% set rf = rifiuti_recupero | default({}) %}

## 1. Dati dell'impianto

- **Denominazione:** {{ impianto.denominazione | default("[DA COMPLETARE: denominazione impianto]") }}
- **Comune:** {{ impianto.comune | default("[DA COMPLETARE: comune]") }} ({{ impianto.provincia | default("[DA COMPLETARE: provincia]") }})

## 2. Inquadramento normativo

{% if rf.pericolosita == "non_pericolosi" %}
L'attività riguarda il recupero di **rifiuti non pericolosi** in procedura
semplificata, secondo le tipologie e le condizioni del **DM 5 febbraio 1998**.
{% elif rf.pericolosita == "pericolosi" %}
L'attività riguarda il recupero di **rifiuti pericolosi** in procedura semplificata,
secondo le tipologie e le condizioni del **DM 12 giugno 2002, n. 161**.
{% else %}
[DA COMPLETARE: pericolosità dei rifiuti — "non_pericolosi" (DM 5/2/1998) oppure
"pericolosi" (DM 161/2002)]
{% endif %}

## 3. Codici EER dei rifiuti gestiti

{% set codici = rf.codici_eer | default([]) %}
{% if codici %}
| Codice EER |
|------------|
{% for c in codici %}
| {{ c }} |
{% endfor %}
{% else %}
[DA COMPLETARE: elenco dei codici EER dei rifiuti gestiti]
{% endif %}

## 4. Operazioni di recupero

{% set operazioni = rf.operazioni | default([]) %}
{% if operazioni %}
Sono effettuate le seguenti operazioni di recupero (Allegato C alla Parte Quarta
del D.Lgs. 152/2006):

{% for op in operazioni %}
- {{ op }}
{% endfor %}
{% else %}
[DA COMPLETARE: operazioni di recupero (codici R, es. R3, R4, R5, R13)]
{% endif %}

## 5. Quantitativi e messa in riserva

- **Quantità annua gestita (t/anno):** {{ rf.quantita_annua_t | default("[DA COMPLETARE: quantità in t/anno]") }}
{% if rf.messa_in_riserva_R13 %}
- È prevista la **messa in riserva (R13)** dei rifiuti in attesa di recupero;
  vanno rispettate le quantità massime stoccabili e le modalità di deposito
  previste dalla normativa di settore.
{% else %}
- Non è prevista la messa in riserva (R13).
{% endif %}

## 6. Modalità di gestione e prescrizioni

[DA COMPLETARE: descrizione delle modalità di stoccaggio, delle aree dedicate,
delle caratteristiche delle materie/prodotti ottenuti dal recupero e del rispetto
delle condizioni tecniche del decreto applicabile]

---

*Relazione redatta a supporto dell'istanza AUA. Tipologie, codici EER, quantità e
prescrizioni vanno verificati rispetto alla normativa e alla modulistica vigenti.*
"""

# Mappa nome-template -> contenuto, usata dal loader Jinja2.
TEMPLATES: dict[str, str] = {
    "modello_unico_aua.md.j2": TEMPLATE_MODELLO_UNICO,
    "relazione_emissioni.md.j2": TEMPLATE_EMISSIONI,
    "relazione_scarichi.md.j2": TEMPLATE_SCARICHI,
    "relazione_rifiuti.md.j2": TEMPLATE_RIFIUTI,
}


# =========================================================================== #
#  CHECKLIST INCORPORATE (vuota da compilare + esempio)                       #
# =========================================================================== #

CHECKLIST_VUOTA = '''# ============ ANAGRAFICA GESTORE ============
gestore:
  ragione_sociale: ""            # OBBLIGATORIO
  forma_giuridica: ""            # es. S.r.l., S.p.A., ditta individuale
  codice_fiscale: ""             # OBBLIGATORIO
  partita_iva: ""
  sede_legale:                   # OBBLIGATORIO
    indirizzo: ""
    comune: ""
    provincia: ""                # sigla, es. CE
    cap: ""
  legale_rappresentante:
    nome: ""
    cognome: ""
    codice_fiscale: ""
    pec: ""                      # necessaria per invio SUAP
  referente_tecnico:
    nome: ""
    cognome: ""
    email: ""
    telefono: ""

# ============ IMPIANTO / INSEDIAMENTO ============
impianto:
  denominazione: ""              # OBBLIGATORIO
  indirizzo: ""
  comune: ""                     # OBBLIGATORIO
  provincia: ""                  # OBBLIGATORIO — atteso: AV/BN/CE/NA/SA
  cap: ""
  foglio: ""                     # dati catastali
  particella: ""
  codice_ateco: ""
  descrizione_attivita: ""       # descrizione sintetica del ciclo produttivo
  superficie_mq:                 # numero
  numero_addetti:                # numero

# ============ TITOLO C/D — EMISSIONI IN ATMOSFERA ============
emissioni_atmosfera:
  presenti: false                # true se ci sono emissioni in atmosfera
  regime: ""                     # "ordinaria" (art.269) | "generale" (art.272 c.2)
  attivita_in_deroga: ""         # se generale: attività dell'Allegato (parte II) richiamato dall'art.272
  punti_emissione:               # elenco dei punti
    - sigla: ""                  # es. E1
      provenienza: ""            # fase/macchinario
      inquinanti: []             # es. [polveri, COV]
      portata_nmc_h:             # numero
      altezza_m:                 # numero
      impianto_abbattimento: ""  # es. filtro a maniche / nessuno

# ============ TITOLO A — SCARICHI DI ACQUE REFLUE ============
scarichi:
  presenti: false
  tipologie: []                  # tra: industriali, domestiche, meteoriche_dilavamento
  recapito: ""                   # "fognatura" | "acque_superficiali" | "suolo"
  corpo_recettore: ""            # denominazione del corpo idrico / gestore fognatura
  portata_mc_anno:               # numero
  presenza_sostanze_pericolose: false   # tab.3/A o 5 All.5 Parte III
  punti_scarico:
    - sigla: ""                  # es. S1
      origine: ""                # processo che genera il refluo
      trattamento: ""            # es. fossa Imhoff / disoleatore / depuratore

# ============ TITOLO G — RECUPERO RIFIUTI (PROCEDURA SEMPLIFICATA) ============
rifiuti_recupero:
  presente: false
  pericolosita: ""               # "non_pericolosi" (DM 5/2/1998) | "pericolosi" (DM 161/2002)
  operazioni: []                 # codici R: es. [R3, R4, R5, R13]
  codici_eer: []                 # es. ["170504", "170302"]
  quantita_annua_t:              # numero (t/anno)
  messa_in_riserva_R13: false    # true se previsto stoccaggio R13

# ============ ALTRI TITOLI (solo rilevazione nell'MVP) ============
altri_titoli:
  utilizzo_agronomico_effluenti: false   # Titolo B (art.112)
  impatto_acustico: false                # Titolo E (L.447/1995)
  fanghi_in_agricoltura: false           # Titolo F (D.Lgs.99/1992)
'''

CHECKLIST_ESEMPIO = '''# ESEMPIO COMPILATO — Officina meccanica con cabina di verniciatura.
# Caso realistico per i test: emissioni in regime generale (art.272 c.2),
# scarichi domestici in fognatura, recupero di rifiuti non pericolosi
# (rottami metallici) con operazioni R13/R4.
# NOTA: dati anagrafici, indirizzi e identificativi sono fittizi.

# ============ ANAGRAFICA GESTORE ============
gestore:
  ragione_sociale: "Autofficina Esposito S.r.l."   # OBBLIGATORIO
  forma_giuridica: "S.r.l."                          # es. S.r.l., S.p.A., ditta individuale
  codice_fiscale: "01234567891"                      # OBBLIGATORIO
  partita_iva: "01234567891"
  sede_legale:                                       # OBBLIGATORIO
    indirizzo: "Via delle Industrie, 24"
    comune: "Caserta"
    provincia: "CE"                                  # sigla, es. CE
    cap: "81100"
  legale_rappresentante:
    nome: "Antonio"
    cognome: "Esposito"
    codice_fiscale: "SPSNTN70A01B963K"
    pec: "autofficina.esposito@pec.it"               # necessaria per invio SUAP
  referente_tecnico:
    nome: "Maria"
    cognome: "Russo"
    email: "ing.russo@studiotecnico.it"
    telefono: "0823 123456"

# ============ IMPIANTO / INSEDIAMENTO ============
impianto:
  denominazione: "Officina meccanica e carrozzeria - sede di Caserta"   # OBBLIGATORIO
  indirizzo: "Via delle Industrie, 24"
  comune: "Caserta"                                  # OBBLIGATORIO
  provincia: "CE"                                    # OBBLIGATORIO — atteso: AV/BN/CE/NA/SA
  cap: "81100"
  foglio: "45"                                       # dati catastali
  particella: "812"
  codice_ateco: "45.20.10"                           # manutenzione e riparazione di autoveicoli
  descrizione_attivita: >-                           # descrizione sintetica del ciclo produttivo
    Officina per la riparazione e manutenzione di autoveicoli con annessa
    carrozzeria. Le lavorazioni comprendono meccanica, lattoneria, stuccatura,
    carteggiatura e verniciatura in cabina-forno pressurizzata. I rottami
    metallici derivanti dalle riparazioni sono stoccati e avviati a recupero.
  superficie_mq: 600                                 # numero
  numero_addetti: 6                                  # numero

# ============ TITOLO C/D — EMISSIONI IN ATMOSFERA ============
emissioni_atmosfera:
  presenti: true                 # true se ci sono emissioni in atmosfera
  regime: "generale"             # "ordinaria" (art.269) | "generale" (art.272 c.2)
  attivita_in_deroga: "Verniciatura, laccatura, doratura di oggetti vari in metallo (cicli a base di prodotti vernicianti)"   # attività dell'Allegato (parte II) richiamato dall'art.272
  punti_emissione:               # elenco dei punti
    - sigla: "E1"                # es. E1
      provenienza: "Cabina-forno di verniciatura"   # fase/macchinario
      inquinanti: [COV, polveri] # es. [polveri, COV]
      portata_nmc_h: 18000       # numero
      altezza_m: 9               # numero
      impianto_abbattimento: "Filtri a secco (pre-filtro su pareti + filtro a tasche sul plenum) e filtri a carboni attivi"  # es. filtro a maniche / nessuno
    - sigla: "E2"
      provenienza: "Postazione di carteggiatura/preparazione"
      inquinanti: [polveri]
      portata_nmc_h: 6000
      altezza_m: 7
      impianto_abbattimento: "Filtro a maniche"

# ============ TITOLO A — SCARICHI DI ACQUE REFLUE ============
scarichi:
  presenti: true
  tipologie: [domestiche]        # tra: industriali, domestiche, meteoriche_dilavamento
  recapito: "fognatura"          # "fognatura" | "acque_superficiali" | "suolo"
  corpo_recettore: "Pubblica fognatura comunale gestita da GORI S.p.A."   # denominazione del corpo idrico / gestore fognatura
  portata_mc_anno: 180           # numero
  presenza_sostanze_pericolose: false   # tab.3/A o 5 All.5 Parte III
  punti_scarico:
    - sigla: "S1"                # es. S1
      origine: "Servizi igienici e spogliatoi del personale"   # processo che genera il refluo
      trattamento: "Fossa Imhoff prima dell'immissione in fognatura"   # es. fossa Imhoff / disoleatore / depuratore

# ============ TITOLO G — RECUPERO RIFIUTI (PROCEDURA SEMPLIFICATA) ============
rifiuti_recupero:
  presente: true
  pericolosita: "non_pericolosi"      # "non_pericolosi" (DM 5/2/1998) | "pericolosi" (DM 161/2002)
  operazioni: [R13, R4]               # codici R: es. [R3, R4, R5, R13]
  codici_eer: ["170405", "160117", "150104"]   # ferro e acciaio; metalli ferrosi; imballaggi metallici
  quantita_annua_t: 40                # numero (t/anno)
  messa_in_riserva_R13: true          # true se previsto stoccaggio R13

# ============ ALTRI TITOLI (solo rilevazione nell'MVP) ============
altri_titoli:
  utilizzo_agronomico_effluenti: false   # Titolo B (art.112)
  impatto_acustico: false                # Titolo E (L.447/1995)
  fanghi_in_agricoltura: false           # Titolo F (D.Lgs.99/1992)
'''


# =========================================================================== #
#  INTAKE — lettura e validazione della checklist                             #
# =========================================================================== #

class ErroreChecklist(Exception):
    """Errore bloccante nella lettura/validazione della checklist."""


SEZIONI_OBBLIGATORIE = ("gestore", "impianto")
CAMPI_GESTORE_OBBLIGATORI = ("ragione_sociale", "codice_fiscale", "sede_legale")
CAMPI_IMPIANTO_OBBLIGATORI = ("denominazione", "comune", "provincia")
PROVINCE_CAMPANIA = {"AV", "BN", "CE", "NA", "SA"}


def carica_checklist(percorso: str | pathlib.Path) -> dict[str, Any]:
    """Carica la checklist YAML e restituisce il dizionario dei dati."""
    percorso = pathlib.Path(percorso)
    if not percorso.exists():
        raise ErroreChecklist(f"File checklist non trovato: {percorso}")
    try:
        with percorso.open("r", encoding="utf-8") as f:
            dati = yaml.safe_load(f)
    except yaml.YAMLError as exc:
        raise ErroreChecklist(f"Errore di sintassi YAML in {percorso}: {exc}") from exc
    if not isinstance(dati, dict):
        raise ErroreChecklist("La checklist deve contenere un dizionario di sezioni.")
    return dati


def valida(dati: dict[str, Any]) -> list[str]:
    """Verifica i requisiti minimi. Restituisce avvisi non bloccanti.

    Solleva ErroreChecklist per problemi che impediscono la generazione.
    """
    avvisi: list[str] = []

    for sezione in SEZIONI_OBBLIGATORIE:
        if sezione not in dati or not isinstance(dati[sezione], dict):
            raise ErroreChecklist(f"Sezione obbligatoria mancante o non valida: '{sezione}'.")

    gestore = dati["gestore"]
    for campo in CAMPI_GESTORE_OBBLIGATORI:
        if not gestore.get(campo):
            raise ErroreChecklist(f"Campo obbligatorio mancante: gestore.{campo}")

    impianto = dati["impianto"]
    for campo in CAMPI_IMPIANTO_OBBLIGATORI:
        if not impianto.get(campo):
            raise ErroreChecklist(f"Campo obbligatorio mancante: impianto.{campo}")

    provincia = str(impianto.get("provincia", "")).upper().strip()
    if provincia not in PROVINCE_CAMPANIA:
        avvisi.append(
            f"La provincia dell'impianto ('{provincia or 'non indicata'}') non risulta tra "
            f"quelle della Regione Campania ({', '.join(sorted(PROVINCE_CAMPANIA))}). "
            "Verificare la competenza territoriale del SUAP/autorità competente."
        )

    if not gestore.get("legale_rappresentante"):
        avvisi.append("Manca il legale rappresentante (gestore.legale_rappresentante): "
                      "necessario per la sottoscrizione dell'istanza.")

    pec = (gestore.get("legale_rappresentante") or {}).get("pec")
    if not pec:
        avvisi.append("Manca la PEC del legale rappresentante: l'invio al SUAP avviene via PEC.")

    return avvisi


# =========================================================================== #
#  APPLICABILITÀ — motore dei 7 titoli AUA (art. 3 DPR 59/2013)               #
# =========================================================================== #

CATALOGO_TITOLI: list[dict[str, Any]] = [
    {"codice": "A", "nome": "Autorizzazione agli scarichi di acque reflue",
     "riferimento": "art. 124 D.Lgs. 152/2006", "implementato": True},
    {"codice": "B", "nome": "Comunicazione per l'utilizzazione agronomica di effluenti, "
     "acque di vegetazione, acque reflue",
     "riferimento": "art. 112 D.Lgs. 152/2006", "implementato": False},
    {"codice": "C", "nome": "Autorizzazione alle emissioni in atmosfera (in via ordinaria)",
     "riferimento": "art. 269 D.Lgs. 152/2006", "implementato": True},
    {"codice": "D", "nome": "Autorizzazione di carattere generale alle emissioni in atmosfera",
     "riferimento": "art. 272 c. 2 D.Lgs. 152/2006", "implementato": True},
    {"codice": "E", "nome": "Comunicazione/documentazione di impatto acustico",
     "riferimento": "L. 447/1995", "implementato": False},
    {"codice": "F", "nome": "Autorizzazione all'utilizzo dei fanghi di depurazione in agricoltura",
     "riferimento": "D.Lgs. 99/1992", "implementato": False},
    {"codice": "G", "nome": "Comunicazioni per il recupero di rifiuti in procedura semplificata",
     "riferimento": "artt. 215-216 D.Lgs. 152/2006", "implementato": True},
]


def _emissioni_regime(dati: dict[str, Any]) -> tuple[bool, str]:
    """Restituisce (presenti, regime) per le emissioni in atmosfera."""
    emissioni = dati.get("emissioni_atmosfera") or {}
    presenti = bool(emissioni.get("presenti"))
    regime = str(emissioni.get("regime") or "").strip().lower()
    return presenti, regime


def _valuta_titolo(codice: str, dati: dict[str, Any]) -> tuple[bool, str]:
    """Applica la regola di applicabilità del singolo titolo."""
    scarichi = dati.get("scarichi") or {}
    rifiuti = dati.get("rifiuti_recupero") or {}
    altri = dati.get("altri_titoli") or {}
    presenti_emissioni, regime = _emissioni_regime(dati)

    if codice == "A":
        if bool(scarichi.get("presenti")):
            return True, "Sono presenti scarichi di acque reflue (scarichi.presenti = true)."
        return False, "Nessuno scarico di acque reflue dichiarato (scarichi.presenti = false)."

    if codice == "B":
        if bool(altri.get("utilizzo_agronomico_effluenti")):
            return True, ("Dichiarata utilizzazione agronomica di effluenti "
                          "(altri_titoli.utilizzo_agronomico_effluenti = true).")
        return False, "Nessuna utilizzazione agronomica di effluenti dichiarata."

    if codice == "C":
        if presenti_emissioni and regime == "ordinaria":
            return True, ("Emissioni in atmosfera in regime ordinario "
                          "(emissioni_atmosfera.regime = 'ordinaria').")
        if presenti_emissioni:
            return False, (f"Emissioni presenti ma in regime '{regime or 'non indicato'}', "
                           "non ordinario.")
        return False, "Nessuna emissione in atmosfera dichiarata."

    if codice == "D":
        if presenti_emissioni and regime == "generale":
            return True, ("Emissioni in atmosfera in regime generale "
                          "(emissioni_atmosfera.regime = 'generale').")
        if presenti_emissioni:
            return False, (f"Emissioni presenti ma in regime '{regime or 'non indicato'}', "
                           "non generale.")
        return False, "Nessuna emissione in atmosfera dichiarata."

    if codice == "E":
        if bool(altri.get("impatto_acustico")):
            return True, "Dichiarato impatto acustico (altri_titoli.impatto_acustico = true)."
        return False, "Nessun impatto acustico dichiarato."

    if codice == "F":
        if bool(altri.get("fanghi_in_agricoltura")):
            return True, ("Dichiarato utilizzo di fanghi in agricoltura "
                          "(altri_titoli.fanghi_in_agricoltura = true).")
        return False, "Nessun utilizzo di fanghi in agricoltura dichiarato."

    if codice == "G":
        if bool(rifiuti.get("presente")):
            return True, ("Previsto recupero di rifiuti in procedura semplificata "
                          "(rifiuti_recupero.presente = true).")
        return False, "Nessun recupero di rifiuti in procedura semplificata dichiarato."

    return False, f"Titolo con codice '{codice}' non riconosciuto."


def valuta_applicabilita(dati: dict[str, Any]) -> list[dict[str, Any]]:
    """Valuta l'applicabilità di tutti i titoli AUA."""
    if not isinstance(dati, dict):
        dati = {}
    esiti: list[dict[str, Any]] = []
    for titolo in CATALOGO_TITOLI:
        codice = titolo["codice"]
        applicabile, motivazione = _valuta_titolo(codice, dati)
        esiti.append({
            "codice": codice,
            "nome": titolo["nome"],
            "riferimento": titolo["riferimento"],
            "applicabile": applicabile,
            "implementato": titolo["implementato"],
            "motivazione": motivazione,
        })
    return esiti


def riepilogo(esiti: list[dict[str, Any]]) -> str:
    """Sintesi testuale degli esiti di applicabilità."""
    da_includere = [e["codice"] for e in esiti if e.get("applicabile")]
    generati = [e["codice"] for e in esiti if e.get("applicabile") and e.get("implementato")]
    da_produrre = [e["codice"] for e in esiti if e.get("applicabile") and not e.get("implementato")]

    def _elenco(codici: list[str]) -> str:
        return ", ".join(codici) if codici else "(nessuno)"

    return (
        f"Titoli da includere: {_elenco(da_includere)} "
        f"— generati: {_elenco(generati)} "
        f"— da produrre a parte: {_elenco(da_produrre)}"
    )


# =========================================================================== #
#  RENDER — Markdown -> DOCX e rendering dei template Jinja2                   #
# =========================================================================== #

_RE_GRASSETTO = re.compile(r"\*\*(.+?)\*\*")

# Ambiente Jinja2 costruito sui template incorporati (DictLoader).
_AMBIENTE_JINJA = Environment(
    loader=DictLoader(TEMPLATES),
    autoescape=select_autoescape(enabled_extensions=(), default=False),
    undefined=ChainableUndefined,
    trim_blocks=True,
    lstrip_blocks=True,
    keep_trailing_newline=True,
)


def _aggiungi_testo_con_grassetto(paragrafo: "Paragraph", testo: str) -> None:
    """Aggiunge testo al paragrafo interpretando il grassetto inline **...**."""
    posizione = 0
    for match in _RE_GRASSETTO.finditer(testo):
        if match.start() > posizione:
            paragrafo.add_run(testo[posizione:match.start()])
        run = paragrafo.add_run(match.group(1))
        run.bold = True
        posizione = match.end()
    if posizione < len(testo):
        paragrafo.add_run(testo[posizione:])


def _e_riga_tabella(riga: str) -> bool:
    spogliata = riga.strip()
    return spogliata.startswith("|") and spogliata.endswith("|")


def _e_riga_separatrice(riga: str) -> bool:
    spogliata = riga.strip().strip("|")
    celle = spogliata.split("|")
    return all(re.fullmatch(r"[\s:-]+", cella) and "-" in cella for cella in celle)


def _celle_riga(riga: str) -> list[str]:
    contenuto = riga.strip().strip("|")
    return [cella.strip() for cella in contenuto.split("|")]


def _aggiungi_tabella(documento: Any, righe_tabella: list[str]) -> None:
    """Costruisce una tabella DOCX da un blocco di righe Markdown."""
    if not righe_tabella:
        return
    matrice = [_celle_riga(riga) for riga in righe_tabella]
    n_colonne = max(len(r) for r in matrice)
    tabella = documento.add_table(rows=0, cols=n_colonne)
    tabella.style = "Table Grid"
    for indice_riga, celle in enumerate(matrice):
        riga_docx = tabella.add_row()
        for indice_colonna in range(n_colonne):
            testo = celle[indice_colonna] if indice_colonna < len(celle) else ""
            paragrafo = riga_docx.cells[indice_colonna].paragraphs[0]
            if indice_riga == 0:
                run = paragrafo.add_run(testo)
                run.bold = True
            else:
                _aggiungi_testo_con_grassetto(paragrafo, testo)


def markdown_to_docx(testo: str, percorso: str) -> None:
    """Converte un sottoinsieme di Markdown in un file .docx."""
    documento = docx.Document()
    righe = testo.splitlines()
    indice = 0
    while indice < len(righe):
        riga = righe[indice]
        spogliata = riga.strip()
        if not spogliata:
            indice += 1
            continue
        if _e_riga_tabella(riga):
            blocco: list[str] = []
            while indice < len(righe) and _e_riga_tabella(righe[indice]):
                if not _e_riga_separatrice(righe[indice]):
                    blocco.append(righe[indice])
                indice += 1
            _aggiungi_tabella(documento, blocco)
            continue
        if spogliata.startswith("### "):
            documento.add_heading(spogliata[4:].strip(), level=3)
            indice += 1
            continue
        if spogliata.startswith("## "):
            documento.add_heading(spogliata[3:].strip(), level=2)
            indice += 1
            continue
        if spogliata.startswith("# "):
            documento.add_heading(spogliata[2:].strip(), level=1)
            indice += 1
            continue
        if spogliata.startswith("- "):
            paragrafo = documento.add_paragraph(style="List Bullet")
            _aggiungi_testo_con_grassetto(paragrafo, spogliata[2:].strip())
            indice += 1
            continue
        paragrafo = documento.add_paragraph()
        _aggiungi_testo_con_grassetto(paragrafo, spogliata)
        indice += 1
    documento.save(percorso)


def renderizza_template(nome_template: str, contesto: dict[str, Any]) -> str:
    """Rende un template incorporato con il contesto dato, restituendo il Markdown."""
    template = _AMBIENTE_JINJA.get_template(nome_template)
    return template.render(**contesto)


# =========================================================================== #
#  CLI — orchestrazione                                                       #
# =========================================================================== #

TEMPLATE_PER_TITOLO: dict[str, dict[str, str]] = {
    "C": {"template": "relazione_emissioni.md.j2", "nome_file": "Relazione_Emissioni"},
    "D": {"template": "relazione_emissioni.md.j2", "nome_file": "Relazione_Emissioni"},
    "A": {"template": "relazione_scarichi.md.j2", "nome_file": "Relazione_Scarichi"},
    "G": {"template": "relazione_rifiuti.md.j2", "nome_file": "Relazione_Rifiuti"},
}
ORDINE_RELAZIONI = ["C", "D", "A", "G"]
_RE_SEGNAPOSTO = re.compile(r"\[DA COMPLETARE[^\]]*\]")


def _costruisci_contesto(dati: dict[str, Any], esiti: list[dict[str, Any]]) -> dict[str, Any]:
    """Costruisce il contesto Jinja2 (sezioni della checklist + lista titoli applicabili)."""
    titoli_applicabili = [
        {"codice": e["codice"], "nome": e["nome"], "riferimento": e["riferimento"]}
        for e in esiti if e.get("applicabile")
    ]
    contesto = dict(dati)
    contesto["titoli"] = titoli_applicabili
    return contesto


def _segnaposto_residui(markdown: str) -> list[str]:
    """Elenco ordinato e senza duplicati dei segnaposto [DA COMPLETARE] in un testo."""
    visti: list[str] = []
    for s in _RE_SEGNAPOSTO.findall(markdown):
        if s not in visti:
            visti.append(s)
    return visti


def _comando_crea_checklist(percorso: str, contenuto: str, descrizione: str) -> int:
    """Scrive un file checklist (vuoto o esempio) sul percorso indicato."""
    destinazione = pathlib.Path(percorso)
    if destinazione.exists():
        print(f"ERRORE: il file esiste già: {destinazione} (non sovrascrivo).", file=sys.stderr)
        return 1
    try:
        if destinazione.parent and not destinazione.parent.exists():
            destinazione.parent.mkdir(parents=True, exist_ok=True)
        destinazione.write_text(contenuto, encoding="utf-8")
    except OSError as exc:
        print(f"ERRORE: impossibile scrivere {destinazione}: {exc}", file=sys.stderr)
        return 1
    print(f"{descrizione} creato in: {destinazione.resolve()}")
    return 0


def _comando_applicabilita(percorso_checklist: str) -> int:
    """Sottocomando 'applicabilita': stampa avvisi e riepilogo dei titoli."""
    try:
        dati = carica_checklist(percorso_checklist)
        avvisi = valida(dati)
    except ErroreChecklist as exc:
        print(f"ERRORE: {exc}", file=sys.stderr)
        return 1

    esiti = valuta_applicabilita(dati)
    if avvisi:
        print("Avvisi:")
        for avviso in avvisi:
            print(f"  - {avviso}")
        print()
    print(riepilogo(esiti))
    print()
    print("Dettaglio titoli:")
    for e in esiti:
        stato = "APPLICABILE" if e["applicabile"] else "non applicabile"
        impl = "generabile" if e["implementato"] else "da produrre a parte"
        print(f"  [{e['codice']}] {e['nome']} — {stato} ({impl})")
        print(f"       {e['motivazione']}")
    return 0


def _comando_genera(percorso_checklist: str, cartella_output: str) -> int:
    """Sottocomando 'genera': produce il pacchetto AUA in cartella_output."""
    try:
        dati = carica_checklist(percorso_checklist)
        avvisi = valida(dati)
    except ErroreChecklist as exc:
        print(f"ERRORE: {exc}", file=sys.stderr)
        return 1

    esiti = valuta_applicabilita(dati)
    contesto = _costruisci_contesto(dati, esiti)
    cartella = pathlib.Path(cartella_output)
    cartella.mkdir(parents=True, exist_ok=True)

    documenti_prodotti: list[str] = []
    segnaposto_per_documento: dict[str, list[str]] = {}

    def _produci(template: str, nome_file: str) -> None:
        try:
            markdown = renderizza_template(template, contesto)
        except TemplateNotFound as exc:
            raise ErroreChecklist(f"Template non trovato: '{template}'.") from exc
        except TemplateError as exc:
            raise ErroreChecklist(f"Errore nel template '{template}': {exc}") from exc
        percorso_docx = cartella / f"{nome_file}.docx"
        try:
            markdown_to_docx(markdown, str(percorso_docx))
        except OSError as exc:
            raise ErroreChecklist(
                f"Impossibile scrivere il documento '{percorso_docx}': {exc}"
            ) from exc
        documenti_prodotti.append(percorso_docx.name)
        residui = _segnaposto_residui(markdown)
        if residui:
            segnaposto_per_documento[percorso_docx.name] = residui

    esiti_per_codice = {e["codice"]: e for e in esiti}
    file_prodotti_set: set[str] = set()
    titoli_da_produrre_a_parte: list[dict[str, Any]] = []

    try:
        _produci("modello_unico_aua.md.j2", "00_Modello_Unico_AUA")
        numero = 1
        for codice in ORDINE_RELAZIONI:
            esito = esiti_per_codice.get(codice)
            if not esito or not esito.get("applicabile"):
                continue
            if not esito.get("implementato"):
                titoli_da_produrre_a_parte.append(esito)
                continue
            info = TEMPLATE_PER_TITOLO.get(codice)
            if info is None:
                titoli_da_produrre_a_parte.append(esito)
                print(
                    f"AVVISO: il titolo {codice} risulta implementato ma manca il "
                    "template associato; verrà elencato tra quelli da produrre a parte.",
                    file=sys.stderr,
                )
                continue
            if info["nome_file"] in file_prodotti_set:
                continue
            _produci(info["template"], f"{numero:02d}_{info['nome_file']}")
            file_prodotti_set.add(info["nome_file"])
            numero += 1
    except ErroreChecklist as exc:
        print(f"ERRORE: {exc}", file=sys.stderr)
        return 1

    for e in esiti:
        if e.get("applicabile") and not e.get("implementato") and e not in titoli_da_produrre_a_parte:
            titoli_da_produrre_a_parte.append(e)

    percorso_checklist_allegati = cartella / "Checklist_Allegati.md"
    _scrivi_checklist_allegati(
        percorso_checklist_allegati,
        documenti_prodotti,
        titoli_da_produrre_a_parte,
        segnaposto_per_documento,
        avvisi,
    )

    if avvisi:
        print("Avvisi:")
        for avviso in avvisi:
            print(f"  - {avviso}")
        print()
    print(riepilogo(esiti))
    print()
    print("Documenti prodotti:")
    for nome in documenti_prodotti:
        print(f"  - {nome}")
    print(f"  - {percorso_checklist_allegati.name}")
    print()
    print(f"Pacchetto generato in: {cartella.resolve()}")
    return 0


def _scrivi_checklist_allegati(
    percorso: pathlib.Path,
    documenti_prodotti: list[str],
    titoli_da_produrre_a_parte: list[dict[str, Any]],
    segnaposto_per_documento: dict[str, list[str]],
    avvisi: list[str],
) -> None:
    """Scrive Checklist_Allegati.md con documenti, titoli mancanti e segnaposto residui."""
    righe: list[str] = []
    righe.append("# Checklist degli allegati e controlli finali — AUA")
    righe.append("")
    righe.append("Documento di servizio generato automaticamente: riepiloga i documenti")
    righe.append("prodotti, i titoli ancora da redigere e i punti `[DA COMPLETARE]` da")
    righe.append("rifinire prima della presentazione al SUAP.")
    righe.append("")
    righe.append("## 1. Documenti prodotti")
    righe.append("")
    for nome in documenti_prodotti:
        righe.append(f"- [ ] {nome}")
    righe.append("")
    righe.append("## 2. Titoli applicabili da produrre a parte")
    righe.append("")
    if titoli_da_produrre_a_parte:
        righe.append("I seguenti titoli risultano applicabili ma non sono ancora generati")
        righe.append("automaticamente dall'MVP: il relativo allegato va prodotto a parte.")
        righe.append("")
        for e in titoli_da_produrre_a_parte:
            righe.append(f"- [ ] **[{e['codice']}] {e['nome']}** ({e['riferimento']})")
    else:
        righe.append("Nessun titolo applicabile resta da produrre a parte.")
    righe.append("")
    righe.append("## 3. Segnaposto `[DA COMPLETARE]` residui")
    righe.append("")
    if segnaposto_per_documento:
        righe.append("Punti da completare manualmente nei documenti prodotti:")
        righe.append("")
        for nome, segnaposto in segnaposto_per_documento.items():
            righe.append(f"### {nome}")
            righe.append("")
            for s in segnaposto:
                righe.append(f"- [ ] {s}")
            righe.append("")
    else:
        righe.append("Nessun segnaposto residuo nei documenti prodotti.")
        righe.append("")
    righe.append("## 4. Avvisi e controlli")
    righe.append("")
    if avvisi:
        for avviso in avvisi:
            righe.append(f"- [ ] {avviso}")
    else:
        righe.append("Nessun avviso dall'intake della checklist.")
    righe.append("")
    percorso.write_text("\n".join(righe), encoding="utf-8")


def _costruisci_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="aua_campania",
        description="Generatore di pratiche AUA (DPR 59/2013) per la Regione Campania — file unico.",
    )
    sottocomandi = parser.add_subparsers(dest="comando", required=True)

    p_crea = sottocomandi.add_parser(
        "crea-checklist", help="Scrive una checklist YAML vuota da compilare.")
    p_crea.add_argument("percorso", nargs="?", default="checklist_AUA.yaml",
                        help="Percorso del file da creare (default: checklist_AUA.yaml).")

    p_es = sottocomandi.add_parser(
        "esempio", help="Scrive una checklist YAML di esempio già compilata.")
    p_es.add_argument("percorso", nargs="?", default="esempio_AUA.yaml",
                      help="Percorso del file da creare (default: esempio_AUA.yaml).")

    p_app = sottocomandi.add_parser(
        "applicabilita", help="Valuta e stampa i titoli AUA applicabili dalla checklist.")
    p_app.add_argument("checklist", help="Percorso del file checklist YAML.")

    p_gen = sottocomandi.add_parser(
        "genera", help="Genera il pacchetto AUA (.docx + Checklist_Allegati.md).")
    p_gen.add_argument("checklist", help="Percorso del file checklist YAML.")
    p_gen.add_argument("-o", "--output", required=True,
                       help="Cartella di destinazione del pacchetto generato.")
    return parser


def main(argv: list[str] | None = None) -> int:
    """Punto d'ingresso della CLI. Restituisce il codice di uscita."""
    parser = _costruisci_parser()
    args = parser.parse_args(argv)

    if args.comando == "crea-checklist":
        return _comando_crea_checklist(args.percorso, CHECKLIST_VUOTA, "Checklist vuota")
    if args.comando == "esempio":
        return _comando_crea_checklist(args.percorso, CHECKLIST_ESEMPIO, "Checklist di esempio")
    if args.comando == "applicabilita":
        return _comando_applicabilita(args.checklist)
    if args.comando == "genera":
        return _comando_genera(args.checklist, args.output)

    parser.print_help()
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
